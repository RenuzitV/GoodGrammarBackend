from docx import Document
from flask import Blueprint, request, jsonify, send_file, abort

import docx
import requests
import io
import re
from bson.binary import Binary
from werkzeug.utils import secure_filename

from middleware.clerk_middleware import token_required
from middleware.stripe_middleware import authentication_and_subscription_threshold_required
from services import user_service
from services.file_service import save_file
from models.file_model import FileObject
from utils.exceptions import UserNotFoundError
from utils.token_utils import get_user_id

bp = Blueprint('file', __name__)

UPLOAD_FOLDER = 'files'
ALLOWED_EXTENSIONS = {'docx'}

API_URL = "https://polite-horribly-cub.ngrok-free.app/generate_code"


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@bp.route('/get_content', methods=['POST'])
def get_file_content():
    # check if the post request has the file part
    if 'file' not in request.files:
        print('No file part')

        return jsonify({'error': 'No file attached'}), 500

    # Get the file
    file = request.files['file']

    # if user does not select file, browser also submit an empty part without filename
    if file.filename == '':
        print('No selected file')

        return jsonify({'error': 'No selected file'}), 500

    # If file valid and is allowed
    if file and allowed_file(file.filename):
        doc = docx.Document(file)

        fulltext = []

        for para in doc.paragraphs:
            fulltext.append(para.text)

        return jsonify({'response': '\n'.join(fulltext)})


@bp.route('/upload', methods=['POST'])
@token_required
# @authentication_and_subscription_threshold_required()
def upload_file(token):
    # Check if the post request has the file part
    if 'file' not in request.files:
        print('No file part')

        return jsonify({'error': 'No file attached'})

    file = request.files['file']

    # if user does not select file, browser also submit an empty part without filename
    if file.filename == '':
        print('No selected file')

        return jsonify({'error': 'No selected file'})

    if file is None or allowed_file(file.filename) is False:
        return jsonify({'error': 'Invalid file type'})

    filename = secure_filename(file.filename)

    resultDoc: Document = docx.Document(file)

    paragraphs = resultDoc.paragraphs

    # Iterate through each paragraph
    for para in paragraphs:

        # Skip the reference
        if is_heading(para) and "reference" in para.text.lower():
            print("Found")
            break

        # If the current paragraph is not heading and not a blank
        if is_heading(para) or para.text.strip() == "":
            continue

        if len(para.runs) == 0 or para.runs[0].text == "" or para.text == "":
            continue

        # Store the first run's formatting
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.bold
        font_italic = first_run.italic
        font_underline = first_run.underline
        font_color = first_run.font.color.rgb

        paragraph_text = ""
        # Aggregate text for the entire paragraph
        for run in para.runs:
            paragraph_text += run.text

        texts = splitKeepDelimiter(paragraph_text)

        # Call the API for language correction on the entire paragraph
        corrected_paragraph = call_API_group(texts)

        # Clear the paragraph and insert the corrected text with the first run's
        para.clear()

        new_run = para.add_run(corrected_paragraph)
        new_run.font.name = font_name
        new_run.font.size = font_size
        new_run.bold = font_bold
        new_run.italic = font_italic
        new_run.underline = font_underline
        if font_color:
            new_run.font.color.rgb = font_color

    # Save the edited document to mongodb
    binaryStream1 = io.BytesIO()

    resultDoc.save(binaryStream1)

    binaryStream1.seek(0)

    encoded = Binary(binaryStream1.read())

    user_id = get_user_id(token)

    myDoc = save_file(filename=filename.split(".")[0] + "-fixed.docx", encodedBinFile=encoded, user_id=user_id)

    try:
        user_service.add_file_to_history(user_id, str(myDoc.id))
    except UserNotFoundError as e:
        print("could not get user", user_id)
        print("Error: ", e)
        abort(404, "User not found")
    except Exception as e:
        print("Error: ", e)
        abort(500, "Internal Server Error")


    return jsonify(
        {
            'edited_file_id': str(myDoc.id)
        }
    )


@bp.route('/get_file_info', methods=['GET'])
def get_file_info():
    fileId = request.args.get('file_id')

    result = FileObject.objects(pk=fileId).first()

    gotFile = result["file"]

    binaryStream2 = io.BytesIO(gotFile)

    myDocument = docx.Document(binaryStream2)

    fulltext = []

    for para in myDocument.paragraphs:
        fulltext.append(para.text)

    return jsonify(
        {
            'file_name': str(result["file_name"]),
            'create_at': str(result["created_at"]),
            'content': '\n'.join(fulltext)
        }
    )


@bp.route('/get_file', methods=['GET'])
def get_file():
    fileId = request.args.get('file_id')

    result = FileObject.objects(pk=fileId).first()

    gotFile = result["file"]

    binaryStream2 = io.BytesIO(gotFile)

    myDocument = docx.Document(binaryStream2)

    f = io.BytesIO()
    myDocument.save(f)
    f.seek(0)

    filename = str(result["file_name"])

    return send_file(
        f,
        as_attachment=True,
        download_name=filename,
    )


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def clean_text(text):
    # Regular expression to find the first occurrence of ".", "!", "?", or newline
    match = re.search(r"[.!?](?=\s|$)|[\n\r]", text)

    if match:
        # Return the substring up to and including the first matched character
        return text[:match.end()]
    else:
        # If no such character is found, return the original text
        return text


def call_API_group(texts):
    # Find the longest string (word-wise) in the list, and multiply by 1.1 to account for any
    # extra characters the promt adds
    max_length = int(max(len(text.split()) for text in texts) * 1.1)
    processed_texts = []
    prompts = []
    param = []
    valid_for_api = []  # check if we should call the API for this text
    # prepend_text = "correct english, no asking, simple answer, only response, keep formatting, no spaces: "
    prepend_text = "Correct English in the following text: "
    # append_text = " Here is the correct version:"
    append_text = ""

    for text in texts:
        # Check if the text contains alphabetic characters
        if re.search("[a-zA-Z]", text) and len(text) < max_length:
            prompts.append(prepend_text + text + append_text)
            valid_for_api.append(True)
        else:
            valid_for_api.append(False)

    if prompts:
        param += [("max_length", max_length)]
        param += [("prompts", prompt) for prompt in prompts]

        response = requests.get(API_URL, params=param)

        if response.status_code == 200:
            body = response.json()
            prompt_index = 0  # To track the index in prompts
            for i, call_api in enumerate(valid_for_api):
                if call_api and prompt_index < len(body):
                    corrected_text = clean_text(body[prompt_index].strip())
                    # print("before: ", texts[i])
                    # print("After: ", corrected_text, "\n\n")
                    processed_texts.append(transform_result(corrected_text, texts[i]))
                    prompt_index += 1
                else:
                    processed_texts.append(texts[i])
        else:
            # print("Failed to retrieve code. Status code:", response.status_code)
            for i, call_api in enumerate(valid_for_api):
                if call_api:
                    processed_texts.append(texts[i])

    return ''.join(processed_texts)


def call_API(text):
    print(str(len(text.split())) + "Before: " + text)

    prompts = []

    result = ""

    prompts.append("Correct English: " + text + "Here is the corrected version:")

    param = []

    param += [("max_length", 128)]

    param += [("prompts", prompt) for prompt in prompts]

    response = requests.get(API_URL, params=param)

    if response.status_code == 200:
        generated_code_list = response.json()
        for i, code in enumerate(generated_code_list):
            print("After: " + transform_result(text, code) + "\n\n")

            return transform_result(text, code)
    else:
        print("Failed to retrieve code. Status code:", response.status_code)

        return "[Failed to correct: " + text


def transform_result(s1, s2):
    result_str = ""

    if s1[0] == ' ':
        result_str = ' '

        if s1[1].islower():
            result_str += s2[0].lower() + s2[1:]

        else:
            result_str += s2[0].upper() + s2[1:]

    else:
        if s1[0].islower():
            result_str = s2[0].lower() + s2[1:]

        else:
            result_str = s2[0].upper() + s2[1:]

    if s1.endswith(' ') and s2.endswith('.'):
        result_str = result_str.rstrip('.')
        result_str += ' '

    return result_str


def is_heading(paragraph):
    """Checks if a paragraph is a heading.

    :param paragraph:
    :return:
    """
    if paragraph.style.name.startswith('Heading'):
        return True


def splitKeepDelimiter(s):
    # Regular expression to split on ".", "!", or "?" but keep the delimiter
    split = re.split(r'([.!?](?=\s|$))', s)

    if len(split) < 2:
        return split

    # Rejoin the split parts to form sentences, excluding newlines
    res = []
    for i in range(0, len(split) - 1, 2):
        combined = split[i] + (split[i + 1] if i + 1 < len(split) else '')
        if combined.strip():
            res.append(combined)

    return res

